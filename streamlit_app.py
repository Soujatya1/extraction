import streamlit as st
import pdfplumber
import pandas as pd
import docx
import io
import os
import zipfile
import openpyxl
from concurrent.futures import ThreadPoolExecutor
import time

def extract_tables_from_pdf(file_path):
    document_content = []
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                document_content.append({
                    "content": text,
                    "page": page_num + 1,
                    "type": "text"
                })
            
            tables = page.extract_tables()
            for table_num, table in enumerate(tables):
                if table:
                    df = pd.DataFrame(table)
                    
                    if not df.empty:
                        headers = []
                        if len(df.columns) > 0:
                            if not pd.isna(df.iloc[0]).all() and not all(x is None for x in df.iloc[0]):
                                headers = [str(h).strip() if h is not None else f"Column_{i}" 
                                          for i, h in enumerate(df.iloc[0])]
                                df = df.iloc[1:]
                            else:
                                headers = [f"Column_{i}" for i in range(len(df.columns))]
                        
                        unique_headers = []
                        header_counts = {}
                        
                        for h in headers:
                            if h in header_counts:
                                header_counts[h] += 1
                                unique_headers.append(f"{h}_{header_counts[h]}")
                            else:
                                header_counts[h] = 0
                                unique_headers.append(h)
                        
                        df.columns = unique_headers
                    
                    document_content.append({
                        "page": page_num + 1,
                        "type": "table",
                        "table_number": table_num + 1,
                        "dataframe": df
                    })
    
    return document_content

def process_single_pdf(uploaded_file, temp_dir):
    """Process a single PDF file and return its content and metadata"""
    file_path = f"{temp_dir}/{uploaded_file.name}"
    
    try:
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        document_content = extract_tables_from_pdf(file_path)
        
        text_count = sum(1 for item in document_content if item["type"] == "text")
        table_count = sum(1 for item in document_content if item["type"] == "table")
        
        return {
            "filename": uploaded_file.name,
            "content": document_content,
            "text_count": text_count,
            "table_count": table_count,
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "filename": uploaded_file.name,
            "content": [],
            "text_count": 0,
            "table_count": 0,
            "success": False,
            "error": str(e)
        }
    
    finally:
        try:
            os.remove(file_path)
        except:
            pass

def create_word_document_text_only(document_content, filename):
    doc = docx.Document()
    doc.add_heading(f'PDF Text Content - {filename}', 0)
    
    text_content = [item for item in document_content if item["type"] == "text"]
    
    for item in text_content:
        doc.add_heading(f'Page {item["page"]}', level=1)
        doc.add_paragraph(item["content"])
    
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io

def create_excel_tables(document_content, filename):
    table_content = [item for item in document_content if item["type"] == "table"]
    
    excel_io = io.BytesIO()
    
    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
        for item in table_content:
            df = item["dataframe"]
            sheet_name = f"Page{item['page']}_Table{item['table_number']}"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            
            counter = 1
            base_sheet_name = sheet_name
            while sheet_name in writer.sheets:
                truncate_length = min(len(base_sheet_name) - len(str(counter)) - 1, 30)
                sheet_name = f"{base_sheet_name[:truncate_length]}_{counter}"
                counter += 1
            
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    excel_io.seek(0)
    return excel_io

def create_combined_zip_archive(all_results):
    """Create a ZIP archive containing all processed files"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for result in all_results:
            if result["success"] and result["content"]:
                base_filename = result["filename"].split('.')[0]
                
                # Add Word document
                word_doc = create_word_document_text_only(result["content"], result["filename"])
                zip_file.writestr(f"{base_filename}_text.docx", word_doc.getvalue())
                
                # Add Excel file if there are tables
                if result["table_count"] > 0:
                    excel_file = create_excel_tables(result["content"], result["filename"])
                    zip_file.writestr(f"{base_filename}_tables.xlsx", excel_file.getvalue())
        
        # Create summary report
        summary_doc = docx.Document()
        summary_doc.add_heading('PDF Processing Summary', 0)
        
        for result in all_results:
            summary_doc.add_heading(result["filename"], level=1)
            if result["success"]:
                summary_doc.add_paragraph(f"✅ Successfully processed")
                summary_doc.add_paragraph(f"Text sections: {result['text_count']}")
                summary_doc.add_paragraph(f"Tables: {result['table_count']}")
            else:
                summary_doc.add_paragraph(f"❌ Processing failed: {result['error']}")
            summary_doc.add_paragraph("")
        
        summary_io = io.BytesIO()
        summary_doc.save(summary_io)
        summary_io.seek(0)
        zip_file.writestr("processing_summary.docx", summary_io.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

st.title("Multi-File PDF Extractor")
st.markdown("Upload multiple PDF files to extract text and tables from all of them at once.")

# File uploader that accepts multiple files
uploaded_files = st.file_uploader(
    "Upload PDF files", 
    type="pdf", 
    accept_multiple_files=True,
    help="You can select multiple PDF files at once"
)

if uploaded_files:
    st.info(f"📁 {len(uploaded_files)} file(s) uploaded")
    
    # Create temporary directory
    temp_dir = 'temp_pdfs'
    os.makedirs(temp_dir, exist_ok=True)
    
    # Process files
    if st.button("Process All Files", type="primary"):
        # Initialize progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        all_results = []
        
        # Process files sequentially with progress updates
        for i, uploaded_file in enumerate(uploaded_files):
            status_text.text(f"Processing {uploaded_file.name}...")
            
            result = process_single_pdf(uploaded_file, temp_dir)
            all_results.append(result)
            
            # Update progress
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        status_text.text("Processing complete!")
        
        st.subheader("Processing Results")
        
        successful_files = [r for r in all_results if r["success"]]
        failed_files = [r for r in all_results if not r["success"]]
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Files", len(uploaded_files))
        with col2:
            st.metric("Successful", len(successful_files), delta=None)
        with col3:
            st.metric("Failed", len(failed_files), delta=None if len(failed_files) == 0 else f"-{len(failed_files)}")
        
        for result in all_results:
            with st.expander(f"📄 {result['filename']}" + ("" if result["success"] else " ❌")):
                if result["success"]:
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**Text sections:** {result['text_count']}")
                    with col2:
                        st.write(f"**Tables:** {result['table_count']}")
                    
                    if result["content"]:
                        text_items = [item for item in result["content"] if item["type"] == "text"]
                        if text_items:
                            st.write("**Text Preview:**")
                            preview_text = text_items[0]["content"][:200] + ("..." if len(text_items[0]["content"]) > 200 else "")
                            st.text(preview_text)
                        
                        table_items = [item for item in result["content"] if item["type"] == "table"]
                        if table_items:
                            st.write("**First Table Preview:**")
                            df = table_items[0]["dataframe"]
                            if not df.empty:
                                st.dataframe(df.head(3))
                else:
                    st.error(f"Error: {result['error']}")
        
        # Download options
        if successful_files:
            st.subheader("Download Options")
            
            # Individual file downloads
            st.write("**Individual Downloads:**")
            for result in successful_files:
                if result["content"]:
                    base_filename = result["filename"].split('.')[0]
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        word_doc = create_word_document_text_only(result["content"], result["filename"])
                        st.download_button(
                            label=f"📄 {base_filename} - Text",
                            data=word_doc,
                            file_name=f"{base_filename}_text.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{base_filename}"
                        )
                    
                    with col2:
                        if result["table_count"] > 0:
                            excel_file = create_excel_tables(result["content"], result["filename"])
                            st.download_button(
                                label=f"📊 {base_filename} - Tables",
                                data=excel_file,
                                file_name=f"{base_filename}_tables.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"excel_{base_filename}"
                            )
                        else:
                            st.write("*No tables found*")
            
            # Combined download
            st.write("**Combined Download:**")
            zip_buffer = create_combined_zip_archive(all_results)
            st.download_button(
                label="📦 Download All Files (ZIP)",
                data=zip_buffer,
                file_name="pdf_extraction_results.zip",
                mime="application/zip",
                help="Downloads all processed files plus a summary report"
            )

else:
    st.info("👆 Please upload one or more PDF files to get started")

try:
    if os.path.exists('temp_pdfs'):
        for file in os.listdir('temp_pdfs'):
            os.remove(os.path.join('temp_pdfs', file))
except:
    pass
